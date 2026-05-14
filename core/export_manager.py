"""
RNV Text Transformer - Export Manager Module
Handles exporting text to multiple formats: TXT, DOCX, HTML, PDF, MD, RTF

Python 3.13 Optimized:
- Modern type hints
- Enum for export formats
- Lazy imports for heavy modules
- Comprehensive error handling

"""

from __future__ import annotations

from enum import StrEnum
from pathlib import Path
from typing import ClassVar
from dataclasses import dataclass
import html

from utils.logger import get_module_logger
from utils.dialog_styles import DialogStyleManager

_logger = get_module_logger("ExportManager")


class ExportFormat(StrEnum):
    """Available export formats."""
    TXT = "Plain Text (.txt)"
    DOCX = "Word Document (.docx)"
    HTML = "HTML Document (.html)"
    PDF = "PDF Document (.pdf)"
    MARKDOWN = "Markdown (.md)"
    RTF = "Rich Text Format (.rtf)"


@dataclass
class ExportOptions:
    """Configuration options for export."""
    format: ExportFormat
    include_metadata: bool = False
    include_line_numbers: bool = False
    preserve_formatting: bool = True
    page_title: str = "Text Transformer Export"
    font_family: str = "Arial"
    font_size: int = 11
    # PDF-specific options
    pdf_page_numbers: bool = True
    pdf_header: str = ""
    pdf_footer: str = ""
    # HTML-specific options
    html_inline_css: bool = True
    html_dark_theme: bool = False


class ExportError(Exception):
    """Custom exception for export errors."""
    __slots__ = ()


class ExportManager:
    """
    Handles exporting text to multiple formats.
    
    Supported formats:
    - TXT: Plain text (UTF-8)
    - DOCX: Microsoft Word document
    - HTML: HTML5 document with optional CSS
    - PDF: PDF document with basic formatting
    - MD: Markdown document
    - RTF: Rich Text Format
    """
    
    # File extension mapping
    EXTENSIONS: ClassVar[dict[ExportFormat, str]] = {
        ExportFormat.TXT: '.txt',
        ExportFormat.DOCX: '.docx',
        ExportFormat.HTML: '.html',
        ExportFormat.PDF: '.pdf',
        ExportFormat.MARKDOWN: '.md',
        ExportFormat.RTF: '.rtf',
    }
    
    # MIME types for each format
    MIME_TYPES: ClassVar[dict[ExportFormat, str]] = {
        ExportFormat.TXT: 'text/plain',
        ExportFormat.DOCX: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        ExportFormat.HTML: 'text/html',
        ExportFormat.PDF: 'application/pdf',
        ExportFormat.MARKDOWN: 'text/markdown',
        ExportFormat.RTF: 'application/rtf',
    }
    
    __slots__ = ()
    
    def __init__(self) -> None:
        """Initialize export manager."""
    
    def export(
        self,
        text: str,
        output_path: str | Path,
        options: ExportOptions
    ) -> bool:
        """
        Export text to the specified format.
        
        Args:
            text: Text content to export
            output_path: Output file path
            options: Export configuration options
            
        Returns:
            True if export successful
            
        Raises:
            ExportError: If export fails
        """
        path = Path(output_path)
        
        # Ensure correct extension
        expected_ext = self.EXTENSIONS[options.format]
        if path.suffix.lower() != expected_ext:
            path = path.with_suffix(expected_ext)
        
        # Ensure parent directory exists
        path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            match options.format:
                case ExportFormat.TXT:
                    self._export_txt(text, path, options)
                case ExportFormat.DOCX:
                    self._export_docx(text, path, options)
                case ExportFormat.HTML:
                    self._export_html(text, path, options)
                case ExportFormat.PDF:
                    self._export_pdf(text, path, options)
                case ExportFormat.MARKDOWN:
                    self._export_markdown(text, path, options)
                case ExportFormat.RTF:
                    self._export_rtf(text, path, options)
                case _:
                    raise ExportError(f"Unsupported format: {options.format}")
            
            if _logger:
                _logger.success(f"Exported to {path.name}", details=options.format.value)
            return True
            
        except ExportError:
            raise
        except Exception as e:
            if _logger:
                _logger.error(f"Export failed: {e}")
            raise ExportError(f"Failed to export as {options.format.value}: {e}") from e
    
    def _export_txt(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as plain text."""
        content = text
        
        if options.include_line_numbers:
            lines = text.splitlines()
            width = len(str(len(lines)))
            content = '\n'.join(
                f"{i+1:>{width}} | {line}" for i, line in enumerate(lines)
            )
        
        path.write_text(content, encoding='utf-8')
    
    def _export_docx(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ExportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Set document properties if metadata enabled
        if options.include_metadata:
            core_props = doc.core_properties
            core_props.title = options.page_title
            core_props.author = "RNV Text Transformer"
        
        # Add title if provided
        if options.page_title and options.page_title != "Text Transformer Export":
            title = doc.add_heading(options.page_title, level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add content
        lines = text.splitlines()
        
        if options.include_line_numbers:
            width = len(str(len(lines)))
            for i, line in enumerate(lines):
                para = doc.add_paragraph(f"{i+1:>{width}} | {line}")
                # Set font
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(options.font_size)
        else:
            for line in lines:
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.name = options.font_family
                    run.font.size = Pt(options.font_size)
        
        doc.save(str(path))
    
    def _export_html(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as HTML document."""
        # Escape HTML entities
        escaped_text = html.escape(text)
        
        # Build line-numbered content if requested
        if options.include_line_numbers:
            lines = escaped_text.splitlines()
            width = len(str(len(lines)))
            line_content = '\n'.join(
                f'<span class="line-num">{i+1:>{width}}</span> {line}'
                for i, line in enumerate(lines)
            )
            content_html = f'<pre class="content">{line_content}</pre>'
        else:
            content_html = f'<pre class="content">{escaped_text}</pre>'
        
        # CSS styles — sourced from DialogStyleManager (no hardcoded hex)
        if options.html_dark_theme:
            _c = DialogStyleManager.DARK
            bg_color = _c['bg']
            text_color = _c['text']
            line_num_color = _c['text_muted']
            accent_color = _c['accent']
        else:
            _c = DialogStyleManager.LIGHT
            bg_color = _c['bg_secondary']
            text_color = _c['text']
            line_num_color = _c['text_muted']
            accent_color = _c['accent']
        
        css = f"""
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: {options.font_family}, sans-serif;
            font-size: {options.font_size}pt;
            line-height: 1.6;
            color: {text_color};
            background-color: {bg_color};
            padding: 20px;
        }}
        h1 {{
            text-align: center;
            color: {accent_color};
            margin-bottom: 20px;
            font-size: 1.5em;
        }}
        .content {{
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
            white-space: pre-wrap;
            word-wrap: break-word;
            background-color: {bg_color};
            padding: 15px;
            border: 1px solid {line_num_color};
            border-radius: 4px;
        }}
        .line-num {{
            color: {line_num_color};
            user-select: none;
            margin-right: 10px;
        }}
        .metadata {{
            font-size: 0.8em;
            color: {line_num_color};
            text-align: center;
            margin-top: 20px;
        }}
        """
        
        # Build HTML document
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="generator" content="RNV Text Transformer">
    <title>{html.escape(options.page_title)}</title>
    {'<style>' + css + '</style>' if options.html_inline_css else ''}
</head>
<body>
    <h1>{html.escape(options.page_title)}</h1>
    {content_html}
    {'<p class="metadata">Generated by RNV Text Transformer</p>' if options.include_metadata else ''}
</body>
</html>
"""
        
        path.write_text(html_content, encoding='utf-8')
    
    def _export_pdf(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as PDF document."""
        try:
            from reportlab.lib.pagesizes import letter, A4  # type: ignore[import-not-found]
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore[import-not-found]
            from reportlab.lib.units import inch  # type: ignore[import-not-found]
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted  # type: ignore[import-not-found]
            from reportlab.lib.enums import TA_LEFT, TA_CENTER  # type: ignore[import-not-found]
        except ImportError:
            raise ExportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(path),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Code'],
            fontName='Courier',
            fontSize=options.font_size,
            leading=options.font_size * 1.4,
            alignment=TA_LEFT
        )
        
        # Build content
        story = []
        
        # Add title
        if options.page_title:
            story.append(Paragraph(options.page_title, title_style))
            story.append(Spacer(1, 12))
        
        # Add content
        lines = text.splitlines()
        
        if options.include_line_numbers:
            width = len(str(len(lines)))
            formatted_lines = [
                f"{i+1:>{width}} | {line}" for i, line in enumerate(lines)
            ]
            content = '\n'.join(formatted_lines)
        else:
            content = text
        
        # Use Preformatted for code-like text
        story.append(Preformatted(content, body_style))
        
        # Add metadata footer
        if options.include_metadata:
            story.append(Spacer(1, 20))
            meta_style = ParagraphStyle(
                'Meta',
                parent=styles['Normal'],
                fontSize=8,
                alignment=TA_CENTER,
                textColor='gray'
            )
            story.append(Paragraph("Generated by RNV Text Transformer", meta_style))
        
        # Build PDF
        doc.build(story)
    
    def _export_markdown(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as Markdown document."""
        lines = []
        
        # Add title
        if options.page_title:
            lines.append(f"# {options.page_title}")
            lines.append("")
        
        # Add content
        if options.include_line_numbers:
            lines.append("```")
            text_lines = text.splitlines()
            width = len(str(len(text_lines)))
            for i, line in enumerate(text_lines):
                lines.append(f"{i+1:>{width}} | {line}")
            lines.append("```")
        else:
            # Wrap in code block to preserve formatting
            lines.append("```")
            lines.append(text)
            lines.append("```")
        
        # Add metadata
        if options.include_metadata:
            lines.append("")
            lines.append("---")
            lines.append("*Generated by RNV Text Transformer*")
        
        path.write_text('\n'.join(lines), encoding='utf-8')
    
    def _export_rtf(self, text: str, path: Path, options: ExportOptions) -> None:
        """Export as Rich Text Format."""
        # RTF header
        rtf_header = r"{\rtf1\ansi\deff0"
        
        # Font table
        rtf_font = r"{\fonttbl{\f0 " + options.font_family + r";}{\f1 Consolas;}}"
        
        # Color table (dark theme colors if applicable)
        rtf_color = r"{\colortbl;\red0\green0\blue0;\red128\green128\blue128;}"
        
        # Document settings
        font_size_twips = options.font_size * 2  # RTF uses half-points
        rtf_settings = f"\\f0\\fs{font_size_twips} "
        
        # Process content
        content_lines = []
        
        # Add title
        if options.page_title:
            title_size = font_size_twips + 8
            content_lines.append(f"\\qc\\b\\fs{title_size} {self._rtf_escape(options.page_title)}\\b0\\par")
            content_lines.append("\\pard\\ql\\par")
        
        # Add text content
        lines = text.splitlines()
        
        if options.include_line_numbers:
            width = len(str(len(lines)))
            for i, line in enumerate(lines):
                num = f"{i+1:>{width}} | "
                escaped = self._rtf_escape(line)
                content_lines.append(f"\\f1\\cf2 {num}\\cf1 {escaped}\\f0\\par")
        else:
            for line in lines:
                escaped = self._rtf_escape(line)
                content_lines.append(f"{escaped}\\par")
        
        # Add metadata
        if options.include_metadata:
            content_lines.append("\\par")
            content_lines.append("\\qc\\fs18\\cf2 Generated by RNV Text Transformer\\par")
        
        # Build RTF document
        rtf_content = rtf_header + rtf_font + rtf_color + rtf_settings + ''.join(content_lines) + "}"
        
        path.write_text(rtf_content, encoding='utf-8')
    
    @staticmethod
    def _rtf_escape(text: str) -> str:
        """Escape special RTF characters."""
        # Escape backslash first
        text = text.replace('\\', '\\\\')
        # Escape braces
        text = text.replace('{', '\\{')
        text = text.replace('}', '\\}')
        # Handle Unicode characters
        result = []
        for char in text:
            code = ord(char)
            if code > 127:
                result.append(f"\\u{code}?")
            else:
                result.append(char)
        return ''.join(result)
    
    @staticmethod
    def get_format_from_extension(extension: str) -> ExportFormat | None:
        """
        Get export format from file extension.
        
        Args:
            extension: File extension (with or without dot)
            
        Returns:
            ExportFormat or None if not recognized
        """
        ext = extension.lower()
        if not ext.startswith('.'):
            ext = '.' + ext
        
        for fmt, fmt_ext in ExportManager.EXTENSIONS.items():
            if fmt_ext == ext:
                return fmt
        return None
    
    @staticmethod
    def get_file_filter() -> str:
        """
        Get file dialog filter string for all supported formats.
        
        Returns:
            Filter string for QFileDialog
        """
        filters = []
        for fmt in ExportFormat:
            ext = ExportManager.EXTENSIONS[fmt]
            filters.append(f"{fmt.value} (*{ext})")
        return ";;".join(filters)
    
    @staticmethod
    def get_available_formats() -> list[ExportFormat]:
        """
        Get list of available export formats.
        
        Returns:
            List of ExportFormat values
        """
        return list(ExportFormat)
    
    @staticmethod
    def check_format_dependencies(fmt: ExportFormat) -> tuple[bool, str]:
        """
        Check if required dependencies are available for a format.
        
        Args:
            fmt: Export format to check
            
        Returns:
            Tuple of (available, message)
        """
        match fmt:
            case ExportFormat.TXT | ExportFormat.HTML | ExportFormat.MARKDOWN:
                return True, "No additional dependencies required"
            
            case ExportFormat.DOCX:
                try:
                    import docx
                    return True, "python-docx is installed"
                except ImportError:
                    return False, "Requires python-docx: pip install python-docx"
            
            case ExportFormat.PDF:
                try:
                    import reportlab  # type: ignore[import-not-found]
                    return True, "reportlab is installed"
                except ImportError:
                    return False, "Requires reportlab: pip install reportlab"
            
            case ExportFormat.RTF:
                return True, "Built-in RTF generation (no dependencies)"
            
            case _:
                return False, "Unknown format"